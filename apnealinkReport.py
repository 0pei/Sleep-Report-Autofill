from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import Select
from selenium.webdriver.support.ui import WebDriverWait
from selenium.common.exceptions import NoSuchElementException
from docx import Document
import time
import sys
import os

def login(chrome):
    account=chrome.find_element(By.ID, "username")
    accountValue=account.get_attribute("value")
    password=chrome.find_element(By.ID, "password")
    passwordValue=password.get_attribute("value")
    if not accountValue:
        account.send_keys(sys.argv[1])
    if not passwordValue:
        password.send_keys(sys.argv[2])
    chrome.find_element(By.ID, "userSubmit").click()

if "crawler" in os.getcwd():
    pythonPath='.\\'
else:
    pythonPath='.\\crawler\\'
service=Service(executable_path=os.path.join(pythonPath, 'chromedriver.exe'))
options=webdriver.ChromeOptions()
chrome=webdriver.Chrome(service=service, options=options)
url="https://ap-airview.resmed.com"
wait=WebDriverWait(chrome, 10)
chrome.get(url)

def is_element_present(by, value):
    try:
        element=chrome.find_element(by, value)
    except NoSuchElementException as e:
        return False
    return True

if(is_element_present(By.ID,"onetrust-accept-btn-handler")==True):
    chrome.find_element(By.ID, "onetrust-accept-btn-handler").click()
    
login(chrome)
time.sleep(10)

diagnostic_href=chrome.find_element(By.XPATH, "//*[@id='diagnostic-patients-link']").get_attribute("href")
chrome.get(diagnostic_href)
time.sleep(1)

select=Select(chrome.find_element(By.ID,"selectPageNum"))
reportValues=[]
dir=sys.argv[3].split('\\')[-1]
for option in select.options:
    if len(reportValues)>1:
        break
    select.select_by_value(str(option.text))
    patients=len(chrome.find_elements(By.XPATH, "//*[@id='hstPatientsTable']/tbody/tr"))
    time.sleep(1)
    for i in range(1, int(patients)+1):
        patientName=chrome.find_element(By.XPATH, "//*[@id='hstPatientsTable']/tbody/tr["+str(i)+"]/td[1]/a").text
        fileNames=patientName.split(', ')
        if len(fileNames[1].split('_'))<2:
            continue
        date=fileNames[1].split('_')[0]
        patientID=fileNames[1].split('_')[1]
        if date==dir.split('_')[0] and patientID==dir.split('_')[1] and fileNames[0]==dir.split('_')[2]:
            patient_href=chrome.find_element(By.XPATH, "//*[@id='hstPatientsTable']/tbody/tr["+str(i)+"]/td[1]/a").get_attribute("href")
            chrome.get(patient_href)
            reportValues=chrome.find_elements(By.CLASS_NAME, "column.report-value")
            break

def replace_text_in_table(table, targetDict):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for targetText, newText in targetDict.items():
                    if targetText in paragraph.text:
                        inline=paragraph.runs
                        inlineText=""
                        for i in range(len(inline)):
                            inlineText=inlineText+inline[i].text
                            inline[i].text=""
                        if targetText in inlineText:
                            text=inlineText.replace(targetText, newText)
                            inline[0].text=text

def main(content):
    if os.path.exists(sys.argv[3]+'\\'+dir+'.docx'):
        try:
            os.remove(sys.argv[3]+'\\'+dir+'.docx')
        except PermissionError:
            pass
        except Exception as e:
            pass
    doc=Document(os.path.join(pythonPath, '成大Home ApneaLink screen報告v2.docx'))
    f=open(sys.argv[3]+'\\reportTemp.txt', 'r')
    lines=f.readlines()
    # allCell=['{a0}', '{a1}', '{a2}', '{a3}', '{a4}', '{a5}', '{a6}', '{a7}', '{a8}', '{a9}', '{a10}', 
    #          '{a11}', '{a12}', '{a13}', '{a14}', '{a15}', '{a16}', '{a17}', '{a18}', '{a19}', '{a20}', 
    #          '{a21}', '{a22}', '{a23}', '{a24}', '{a25}', 
    #          '{b1}', '{b2}', '{b3}', '{b4}', '{b5}', '{b6}', '{b7}', '{b8}', '{b9}', '{b10}', 
    #          '{b11}', '{b12}', '{b13}', '{b14}', '{b15}', '{b16}', 
    #          '{c1}', '{c2}', '{c3}', '{c4}', '{c5}', '{c6}', '{c7}', '{c8}', '{c9}', '{c10}',
    #          '{c11}', '{c12}', '{c13}', 
    #          '{d1}', '{d2}', '{d3}', '{d4}', '{d5}', '{d6}', '{d7}', '{d8}', '{d9}', '{d10}',
    #          '{e1}', '{e2}', '{e3}', '{e4}', 
    #          '{f1}', '{f2}', 
    #          '{g1}', '{g2}', '{g3}', '{g4}', '{g5}', '{g6}', '{g7}', '{g8}', '{g9}', '{g10}', '{g11}',
    #          '{h1}', '{h2}', '{h3}', '{h4}', '{h5}', '{h6}', '{h7}', '{h8}', '{h9}'] 
    canFillCell=['{a0}', '{a3}', '{a1}', 
                 '{b1}', '{b2}', '{b3}', '{b4}', '{b5}', '{b6}', '{b7}', '{b8}', '{b9}', '{b10}', 
                 '{b11}', '{b12}', '{b13}', '{b14}', '{b15}', '{b16}', 
                 '{c1}', '{c2}', '{c3}', '{c4}', '{c5}', '{c6}', '{c7}', '{c8}', '{c9}', '{c10}',
                 '{c11}', '{c12}', '{c13}',
                 '{d1}', '{d2}', '{d3}', '{d4}', '{d5}', '{d6}', '{d7}', '{d8}', '{d9}', '{d10}',
                 '{e1}', '{e2}', '{e3}', '{e4}',
                 '{g1}', '{g2}', '{g3}', '{g4}', '{g5}', '{g6}', '{g7}', '{g8}', '{g9}', '{g10}', '{g11}',]
    replacementDict={
        # cell: content[i] for i, cell in enumerate(canFillCell)
        '{a1}': lines[0].strip(), #reportValues[13].text,
        '{a2}': lines[1].strip(), #reportValues[12].text,
        '{a3}': lines[2].strip(), #reportValues[11].text,
        '{a4}': lines[3].strip(), #reportValues[30].text,
        '{a5}': lines[4].strip(), #reportValues[29].text,
        '{a6}': lines[5].strip(), #reportValues[34].text,
        '{a7}': lines[6].strip(), #reportValues[33].text,
        '{a8}': lines[7].strip(), #reportValues[32].text,
        '{a9}': lines[8].strip(), #reportValues[31].text,
        # '{a10}': reportValues[16].text,
        # '{a11}': reportValues[15].text,
        # '{a12}': reportValues[14].text,
        # '{a13}': reportValues[18].text,
        # '{a14}': reportValues[17].text[1:-2],
        # '{a15}': reportValues[21].text,
        # '{a16}': reportValues[20].text,
        # '{a17}': reportValues[19].text,
        # '{a18}': reportValues[23].text,
        # '{a19}': reportValues[22].text[1:-2],
        # '{a20}': reportValues[26].text,
        # '{a21}': reportValues[25].text,
        # '{a22}': reportValues[24].text,
        # '{a23}': reportValues[28].text,
        # '{a24}': reportValues[27].text[1:-2],
        '{a25}': lines[24].strip(), #reportValues[36].text,
        '{a26}': lines[25].strip(), #reportValues[35].text,

        '{b1}': reportValues[3].text,
        '{b2}': reportValues[2].text,
        '{b3}': reportValues[1].text,
        '{b4}': reportValues[7].text,
        '{b5}': reportValues[6].text,
        '{b6}': reportValues[5].text,
        '{b7}': reportValues[10].text,
        '{b8}': reportValues[9].text,
        '{b9}': reportValues[8].text,

        '{c1}': reportValues[41].text,
        '{c2}': lines[36].strip(), #reportValues[40].text,
        '{c3}': lines[37].strip(), #reportValues[39].text,
        '{c4}': reportValues[44].text,
        '{c5}': reportValues[43].text,
        '{c6}': reportValues[42].text,
        '{c7}': reportValues[46].text,
        '{c8}': reportValues[45].text,
        '{c9}': lines[43].strip(), #reportValues[37].text,
        '{c10}': lines[44].strip(), #reportValues[38].text,

        '{d1}': reportValues[50].text,
        '{d2}': reportValues[49].text,
        '{d3}': reportValues[48].text,

        '{e1}': lines[48].strip(), #reportValues[53].text,
        '{e2}': lines[49].strip(), #reportValues[52].text,
        '{e3}': lines[50].strip(), #reportValues[51].text,
    }
    replace_text_in_table(doc.tables[0], replacementDict)
    replace_text_in_table(doc.tables[2], replacementDict)
    doc.save(sys.argv[3]+'\\'+dir+'.docx')
    f.close()
    os.remove(sys.argv[3]+'\\reportTemp.txt')

if __name__ == "__main__":
    main()